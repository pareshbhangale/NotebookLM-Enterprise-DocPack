from __future__ import annotations

"""NotebookLM Enterprise DocPack: turn enterprise documentation into upload-ready DOCX packs.

Flow:
1. Read a manifest CSV with source URLs.
2. Download source files into a raw folder.
3. Convert every source file to Markdown with MarkItDown.
4. Merge all Markdown files into one master Markdown file.
5. Build source-labeled DOCX packs that stay below the target word count.

Default project layout:
- /Input for manifests and other source inputs
- /Output for generated pipeline artifacts

Use this script as the project's command-line entry point.
"""

import argparse
import csv
import os
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath
from urllib.parse import unquote, urlparse

try:
    import requests
except Exception:
    requests = None  # type: ignore

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None  # type: ignore
    Pt = None  # type: ignore

try:
    from markitdown import MarkItDown
except Exception:
    MarkItDown = None  # type: ignore


PROJECT_ROOT = Path(__file__).resolve().parent


def project_dir(preferred_name: str, fallback_name: str) -> Path:
    """Prefer the checked-in folder casing, but tolerate older lowercase clones."""
    preferred = PROJECT_ROOT / preferred_name
    fallback = PROJECT_ROOT / fallback_name
    if preferred.exists() or not fallback.exists():
        return preferred
    return fallback


INPUT_DIR = project_dir("Input", "input")
OUTPUT_DIR = project_dir("Output", "output")

DEFAULT_MANIFEST = INPUT_DIR / "manifest.csv"
EXAMPLE_MANIFEST = INPUT_DIR / "example_download_manifest.csv"
DEFAULT_RAW = OUTPUT_DIR / "raw"
DEFAULT_MD = OUTPUT_DIR / "markdown"
DEFAULT_MASTER_MD = OUTPUT_DIR / "knowledge_pack.md"
DEFAULT_DOCX = OUTPUT_DIR / "docx_packs"
DEFAULT_SUMMARY = OUTPUT_DIR / "pipeline_summary.txt"

DEFAULT_CHUNK_WORDS = 75_000
DEFAULT_PACK_WORDS = 375_000
DEFAULT_HARD_MAX_WORDS = 400_000
DEFAULT_MAX_DOCX_PARTS = 0
DOCX_CODE_FONT = "Menlo" if sys.platform == "darwin" else "Consolas"

WORD_RE = re.compile(r"\S+")
FILE_BLOCK_RE = re.compile(
    r"(?ms)^---\s*\n# FILE:\s*(.*?)\s*\n---\s*\n(.*?)(?=^---\s*\n# FILE:\s*|\Z)"
)


@dataclass
class DownloadStats:
    downloaded: int = 0
    skipped_existing: int = 0
    failed: int = 0


@dataclass
class ConvertStats:
    converted: int = 0
    skipped_existing: int = 0
    failed: int = 0


def count_words(text: str) -> int:
    return len(WORD_RE.findall(text))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def install_hint() -> str:
    return "Run: ./setup.sh or python3 -m pip install -r requirements.txt"


def normalize_cli_path(path_value: str | Path) -> str | Path:
    if isinstance(path_value, str) and os.sep == "/":
        return path_value.replace("\\", "/")
    return path_value


def project_relative_path(path_value: str | Path) -> Path:
    relative = Path(normalize_cli_path(path_value))
    parts = relative.parts
    if not parts:
        return PROJECT_ROOT

    for child in PROJECT_ROOT.iterdir():
        if child.name.lower() == parts[0].lower():
            return child.joinpath(*parts[1:])
    return PROJECT_ROOT / relative


def remove_generated_path(path: Path, label: str) -> None:
    if not path.exists():
        return

    if path.is_dir():
        shutil.rmtree(path)
    else:
        path.unlink()
    print(f"[clean] removed {label}: {path}")


def clean_output_for_run(
    raw_dir: Path,
    md_dir: Path,
    master_md: Path,
    docx_dir: Path,
    summary_file: Path,
    skip_download: bool,
    skip_convert: bool,
    skip_merge: bool,
    skip_docx: bool,
) -> None:
    """Remove generated artifacts for stages that will run in this invocation."""
    remove_generated_path(summary_file, "summary")

    if not skip_download:
        remove_generated_path(raw_dir, "raw downloads")
        remove_generated_path(md_dir, "markdown")
        remove_generated_path(master_md, "merged markdown")
        remove_generated_path(docx_dir, "docx packs")
        return

    if not skip_convert:
        remove_generated_path(md_dir, "markdown")
        remove_generated_path(master_md, "merged markdown")
        remove_generated_path(docx_dir, "docx packs")
        return

    if not skip_merge:
        remove_generated_path(master_md, "merged markdown")
        remove_generated_path(docx_dir, "docx packs")
        return

    if not skip_docx:
        remove_generated_path(docx_dir, "docx packs")


def resolve_project_path(path_value: str | Path, label: str) -> Path:
    normalized_path = normalize_cli_path(path_value)
    if isinstance(normalized_path, str) and os.sep == "/" and re.match(r"^[A-Za-z]:/", normalized_path):
        raise ValueError(f"{label} uses a Windows drive path. Use a macOS path instead.")

    candidate = Path(normalized_path)
    if not candidate.is_absolute():
        candidate = project_relative_path(candidate)

    resolved = candidate.resolve(strict=False)
    project_root = PROJECT_ROOT.resolve()
    resolved_norm = os.path.normcase(str(resolved))
    project_norm = os.path.normcase(str(project_root))
    if os.path.commonpath([resolved_norm, project_norm]) != project_norm:
        raise ValueError(f"{label} must stay inside the project folder: {project_root}")

    return resolved


def normalize_manifest_row(row: dict[str, str | None]) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for key, value in row.items():
        clean_key = (key or "").strip().lower()
        normalized[clean_key] = (value or "").strip()
    return normalized


def read_manifest_csv(manifest_path: Path) -> list[tuple[str, str, str]]:
    if not manifest_path.exists():
        if EXAMPLE_MANIFEST.exists():
            raise FileNotFoundError(
                f"Manifest CSV not found: {manifest_path}. "
                f"Copy {EXAMPLE_MANIFEST.name} to {manifest_path.name} and edit it."
            )
        raise FileNotFoundError(f"Manifest CSV not found: {manifest_path}")

    rows: list[tuple[str, str, str]] = []
    with manifest_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for raw_row in reader:
            row = normalize_manifest_row(raw_row)
            url = row.get("url") or row.get("link") or row.get("source_url") or ""
            if not url:
                continue
            kind = row.get("kind") or row.get("type") or "document"
            title = row.get("title") or row.get("name") or filename_from_url(url)
            rows.append((kind, title, url))

    if not rows:
        raise RuntimeError(f"No URLs found in manifest: {manifest_path}")
    return rows


def filename_from_url(url: str) -> str:
    parsed = urlparse(url)
    name = PurePosixPath(parsed.path.replace("\\", "/")).name
    name = unquote(name)
    return name or "downloaded_file"


def download_file(url: str, out_path: Path, timeout: int = 120) -> str:
    if requests is None:
        raise RuntimeError(f"requests is not installed. {install_hint()}")

    ensure_dir(out_path.parent)
    if out_path.exists() and out_path.stat().st_size > 0:
        print(f"[download] skip existing: {out_path.name}")
        return "skipped"

    print(f"[download] {out_path.name}")
    with requests.get(url, stream=True, timeout=timeout) as response:
        response.raise_for_status()
        with out_path.open("wb") as handle:
            for chunk in response.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    handle.write(chunk)
    return "downloaded"


def download_all(manifest_path: Path, raw_dir: Path) -> DownloadStats:
    ensure_dir(raw_dir)
    stats = DownloadStats()
    for _, _, url in read_manifest_csv(manifest_path):
        out_path = raw_dir / filename_from_url(url)
        try:
            result = download_file(url, out_path)
            if result == "downloaded":
                stats.downloaded += 1
            else:
                stats.skipped_existing += 1
        except Exception as exc:
            stats.failed += 1
            print(f"[WARN] download failed: {url} -> {exc}")
    return stats


def unpack_zip(zip_path: Path, extract_root: Path) -> Path:
    target = extract_root / zip_path.stem
    if target.exists() and any(target.iterdir()):
        return target

    ensure_dir(target)
    try:
        with zipfile.ZipFile(zip_path, "r") as archive:
            extract_zip_safely(archive, target)
        print(f"[zip] extracted: {zip_path.name}")
    except Exception as exc:
        print(f"[WARN] zip extract failed: {zip_path} -> {exc}")
    return target


def safe_zip_member_parts(member_name: str) -> list[str] | None:
    normalized = member_name.replace("\\", "/")
    parts = [part for part in PurePosixPath(normalized).parts if part not in {"", ".", "/"}]
    if not parts or ".." in parts or parts[0].endswith(":"):
        return None
    return parts


def extract_zip_safely(archive: zipfile.ZipFile, target: Path) -> None:
    target_root = target.resolve(strict=False)
    for member in archive.infolist():
        parts = safe_zip_member_parts(member.filename)
        if parts is None:
            print(f"[WARN] zip skipped unsafe member: {member.filename}")
            continue

        out_path = target.joinpath(*parts)
        out_resolved = out_path.resolve(strict=False)
        if os.path.commonpath([str(out_resolved), str(target_root)]) != str(target_root):
            print(f"[WARN] zip skipped outside member: {member.filename}")
            continue

        if member.is_dir():
            ensure_dir(out_path)
            continue

        ensure_dir(out_path.parent)
        with archive.open(member, "r") as source, out_path.open("wb") as dest:
            shutil.copyfileobj(source, dest)


def iter_source_files(raw_dir: Path, extracted_root: Path):
    for path in raw_dir.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(raw_dir)
        if rel.parts and rel.parts[0] == extracted_root.name:
            continue
        if path.name.startswith("~$") or path.name.startswith(".~"):
            continue
        if path.suffix.lower() == ".zip":
            extracted = unpack_zip(path, extracted_root)
            for child in extracted.rglob("*"):
                if child.is_file() and not child.name.startswith(("~$", ".~")):
                    yield child
        else:
            yield path


def markitdown_convert(path: Path) -> str:
    if MarkItDown is None:
        raise RuntimeError(
            f"markitdown is not installed. {install_hint()}"
        )

    try:
        converter = MarkItDown(enable_plugins=True)
    except TypeError:
        converter = MarkItDown()
    result = converter.convert(str(path))
    text = getattr(result, "text_content", None)
    return (text if text is not None else str(result)).strip()


def convert_to_md(raw_dir: Path, md_dir: Path) -> ConvertStats:
    if not raw_dir.exists():
        raise FileNotFoundError(f"Raw input folder not found: {raw_dir}")
    if MarkItDown is None:
        raise RuntimeError(f"markitdown is not installed. {install_hint()}")

    ensure_dir(md_dir)
    extracted_root = raw_dir / "_unzipped"
    ensure_dir(extracted_root)
    stats = ConvertStats()

    for src in iter_source_files(raw_dir, extracted_root):
        rel = src.relative_to(raw_dir)

        out_path = md_dir / rel.with_suffix(".md")
        ensure_dir(out_path.parent)
        if out_path.exists() and out_path.stat().st_size > 0:
            stats.skipped_existing += 1
            continue

        try:
            md_text = markitdown_convert(src)
            if md_text.strip():
                out_path.write_text(md_text + "\n", encoding="utf-8")
                stats.converted += 1
                print(f"[md] {out_path.relative_to(md_dir)}")
        except Exception as exc:
            stats.failed += 1
            print(f"[WARN] convert failed: {src} -> {exc}")

    return stats


def merge_md_files(md_dir: Path, master_md: Path) -> int:
    if not md_dir.exists():
        raise FileNotFoundError(f"Markdown folder not found: {md_dir}")

    ensure_dir(master_md.parent)
    md_files = sorted(md_dir.rglob("*.md"), key=lambda path: str(path).lower())
    if not md_files:
        raise RuntimeError(f"No .md files found in {md_dir}")

    written_sources = 0
    with master_md.open("w", encoding="utf-8") as out:
        for md_file in md_files:
            rel = md_file.relative_to(md_dir).as_posix()
            text = md_file.read_text(encoding="utf-8", errors="ignore").strip()
            if not text:
                continue
            out.write("\n\n---\n")
            out.write(f"# FILE: {rel}\n")
            out.write("---\n\n")
            out.write(text)
            out.write("\n")
            written_sources += 1

    print(f"[merge] wrote {master_md}")
    return written_sources


def parse_master_md(md_text: str) -> list[tuple[str, str]]:
    blocks = FILE_BLOCK_RE.findall(md_text)
    if blocks:
        return [(src.strip(), body.strip()) for src, body in blocks if body.strip()]

    text = md_text.strip()
    return [("MASTER", text)] if text else []


def split_text_by_words(text: str, max_words: int) -> list[str]:
    if count_words(text) <= max_words:
        return [text.strip()]

    paras = [para.strip() for para in re.split(r"\n\s*\n", text) if para.strip()]
    chunks: list[str] = []
    buf: list[str] = []
    buf_words = 0

    def flush() -> None:
        nonlocal buf, buf_words
        if buf:
            chunks.append("\n\n".join(buf).strip())
            buf = []
            buf_words = 0

    for para in paras:
        para_words = count_words(para)
        if para_words > max_words:
            flush()
            sentences = re.split(r"(?<=[.!?])\s+", para)
            sent_buf: list[str] = []
            sent_words = 0
            for sent in sentences:
                sent = sent.strip()
                if not sent:
                    continue
                sentence_words = count_words(sent)
                if sentence_words > max_words:
                    words = WORD_RE.findall(sent)
                    for index in range(0, len(words), max_words):
                        chunks.append(" ".join(words[index : index + max_words]))
                    continue
                if sent_words + sentence_words > max_words and sent_buf:
                    chunks.append(" ".join(sent_buf).strip())
                    sent_buf = []
                    sent_words = 0
                sent_buf.append(sent)
                sent_words += sentence_words
            if sent_buf:
                chunks.append(" ".join(sent_buf).strip())
            continue

        if buf and buf_words + para_words > max_words:
            flush()
        buf.append(para)
        buf_words += para_words

    flush()
    return [chunk for chunk in chunks if chunk.strip()]


def markdown_to_docx(md_text: str, doc: Document) -> None:
    if Pt is None:
        raise RuntimeError(f"python-docx is not installed. {install_hint()}")

    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    in_code = False
    code_buf: list[str] = []

    def add_code(block_lines: list[str]) -> None:
        if not block_lines:
            return
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("\n".join(block_lines))
        run.font.name = DOCX_CODE_FONT
        run.font.size = Pt(9)

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code(code_buf)
                code_buf = []
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not line.strip():
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            doc.add_heading(
                heading_match.group(2).strip(),
                level=min(len(heading_match.group(1)), 4),
            )
            continue
        if re.match(r"^[-*+]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*+]\s+", "", line).strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
            continue
        if line.strip() in {"---", "***", "___"}:
            doc.add_paragraph()
            continue
        doc.add_paragraph(line.strip())

    if in_code and code_buf:
        add_code(code_buf)


def build_docx_pack(
    blocks: list[tuple[str, str]],
    out_dir: Path,
    master_md: Path,
    chunk_words: int,
    pack_words: int,
    max_words: int,
    max_docx_parts: int,
) -> list[Path]:
    if Document is None:
        raise RuntimeError(f"python-docx is not installed. {install_hint()}")

    ensure_dir(out_dir)
    written: list[Path] = []
    units: list[tuple[str, str, int]] = []

    for src, body in blocks:
        pieces = split_text_by_words(body, chunk_words)
        for index, piece in enumerate(pieces, start=1):
            title = f"{src} ({index}/{len(pieces)})"
            units.append((title, piece, count_words(piece)))

    def count_parts_for_capacity(capacity: int) -> int:
        parts = 0
        current_count = 0
        for _, _, word_count in units:
            if current_count and current_count + word_count > capacity:
                parts += 1
                current_count = 0
            current_count += word_count
        if current_count:
            parts += 1
        return parts

    if max_docx_parts > 0:
        if count_parts_for_capacity(max_words) > max_docx_parts:
            raise RuntimeError(
                "Unable to keep DOCX output within "
                f"{max_docx_parts} files under {max_words} words each. "
                "Lower --chunk-words or raise --max-words."
            )

        if count_parts_for_capacity(pack_words) > max_docx_parts:
            low = pack_words + 1
            high = max_words
            while low < high:
                midpoint = (low + high) // 2
                if count_parts_for_capacity(midpoint) <= max_docx_parts:
                    high = midpoint
                else:
                    low = midpoint + 1
            print(f"[docx] adjusted pack words from {pack_words} to {low}")
            pack_words = low

    part_idx = 1
    current: list[tuple[str, str, int]] = []
    current_words = 0

    def flush() -> None:
        nonlocal part_idx, current, current_words
        if not current:
            return

        out_file = out_dir / f"knowledge_pack_part_{part_idx:03d}.docx"
        doc = Document()
        doc.add_heading(f"Knowledge Pack - Part {part_idx}", level=0)
        doc.add_paragraph(f"Source: {master_md}")
        doc.add_paragraph("")

        for title, text, _ in current:
            doc.add_heading(title, level=1)
            markdown_to_docx(text, doc)
            doc.add_paragraph("")

        doc.save(str(out_file))
        check_doc = Document(str(out_file))
        word_count = count_words(" ".join(p.text for p in check_doc.paragraphs if p.text.strip()))
        print(f"[docx] wrote {out_file.name} | words={word_count}")
        if word_count > max_words:
            print(f"[WARN] {out_file.name} exceeds {max_words} words.")

        written.append(out_file)
        part_idx += 1
        current = []
        current_words = 0

    for title, text, word_count in units:
        if current and current_words + word_count > pack_words:
            flush()
        current.append((title, text, word_count))
        current_words += word_count

    flush()
    if max_docx_parts > 0 and len(written) > max_docx_parts:
        raise RuntimeError(
            f"Generated {len(written)} DOCX files, exceeding --max-docx-parts "
            f"{max_docx_parts}."
        )
    return written


def write_summary(
    summary_path: Path,
    manifest: Path,
    raw_dir: Path,
    md_dir: Path,
    master_md: Path,
    docx_dir: Path,
    download_stats: DownloadStats | None,
    convert_stats: ConvertStats | None,
    merged_sources: int | None,
    docx_files: list[Path] | None,
) -> None:
    ensure_dir(summary_path.parent)
    lines = [
        "NotebookLM Enterprise DocPack pipeline summary",
        f"manifest={manifest}",
        f"raw_dir={raw_dir}",
        f"markdown_dir={md_dir}",
        f"master_markdown={master_md}",
        f"docx_dir={docx_dir}",
    ]

    if download_stats is not None:
        lines.extend(
            [
                f"downloaded={download_stats.downloaded}",
                f"download_skipped_existing={download_stats.skipped_existing}",
                f"download_failed={download_stats.failed}",
            ]
        )
    if convert_stats is not None:
        lines.extend(
            [
                f"markdown_converted={convert_stats.converted}",
                f"markdown_skipped_existing={convert_stats.skipped_existing}",
                f"markdown_failed={convert_stats.failed}",
            ]
        )
    if merged_sources is not None:
        lines.append(f"merged_sources={merged_sources}")
    if docx_files is not None:
        lines.append(f"docx_parts={len(docx_files)}")

    summary_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"[summary] wrote {summary_path}")


def validate_args(args: argparse.Namespace) -> None:
    if args.chunk_words <= 0:
        raise ValueError("--chunk-words must be greater than zero.")
    if args.pack_words <= 0:
        raise ValueError("--pack-words must be greater than zero.")
    if args.max_words <= 0:
        raise ValueError("--max-words must be greater than zero.")
    if args.max_docx_parts < 0:
        raise ValueError("--max-docx-parts cannot be negative.")
    if args.chunk_words > args.max_words:
        raise ValueError("--chunk-words cannot be greater than --max-words.")
    if args.pack_words > args.max_words:
        raise ValueError("--pack-words cannot be greater than --max-words.")


def validate_stage_inputs(
    manifest: Path,
    raw_dir: Path,
    md_dir: Path,
    master_md: Path,
    skip_download: bool,
    skip_convert: bool,
    skip_merge: bool,
    skip_docx: bool,
) -> None:
    """Check inputs needed by skipped stages before clearing any output."""
    if not skip_download:
        read_manifest_csv(manifest)
    elif not skip_convert and not raw_dir.exists():
        raise FileNotFoundError(
            f"Raw input folder not found: {raw_dir}. "
            "Run without --skip-download or provide an existing --raw-dir."
        )

    if skip_convert and not skip_merge and not md_dir.exists():
        raise FileNotFoundError(
            f"Markdown folder not found: {md_dir}. "
            "Run without --skip-convert or provide an existing --md-dir."
        )

    if skip_merge and not skip_docx and not master_md.exists():
        raise FileNotFoundError(
            f"Master Markdown file not found: {master_md}. "
            "Run without --skip-merge or provide an existing --master-md."
        )


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="notebooklm_enterprise_docpack.py",
        description=(
            "NotebookLM Enterprise DocPack: download documents -> Markdown -> merged Markdown -> "
            "upload-ready DOCX packs. All runtime paths stay inside this project folder."
        ),
    )
    parser.add_argument("--manifest", default=str(DEFAULT_MANIFEST))
    parser.add_argument("--raw-dir", default=str(DEFAULT_RAW))
    parser.add_argument("--md-dir", default=str(DEFAULT_MD))
    parser.add_argument("--master-md", default=str(DEFAULT_MASTER_MD))
    parser.add_argument("--docx-dir", default=str(DEFAULT_DOCX))
    parser.add_argument("--summary-file", default=str(DEFAULT_SUMMARY))
    parser.add_argument("--chunk-words", type=int, default=DEFAULT_CHUNK_WORDS)
    parser.add_argument("--pack-words", type=int, default=DEFAULT_PACK_WORDS)
    parser.add_argument("--max-words", type=int, default=DEFAULT_HARD_MAX_WORDS)
    parser.add_argument("--max-docx-parts", type=int, default=DEFAULT_MAX_DOCX_PARTS)
    parser.add_argument("--skip-download", action="store_true")
    parser.add_argument("--skip-convert", action="store_true")
    parser.add_argument("--skip-merge", action="store_true")
    parser.add_argument("--skip-docx", action="store_true")
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    validate_args(args)

    manifest = resolve_project_path(args.manifest, "--manifest")
    raw_dir = resolve_project_path(args.raw_dir, "--raw-dir")
    md_dir = resolve_project_path(args.md_dir, "--md-dir")
    master_md = resolve_project_path(args.master_md, "--master-md")
    docx_dir = resolve_project_path(args.docx_dir, "--docx-dir")
    summary_file = resolve_project_path(args.summary_file, "--summary-file")

    validate_stage_inputs(
        manifest=manifest,
        raw_dir=raw_dir,
        md_dir=md_dir,
        master_md=master_md,
        skip_download=args.skip_download,
        skip_convert=args.skip_convert,
        skip_merge=args.skip_merge,
        skip_docx=args.skip_docx,
    )

    ensure_dir(INPUT_DIR)
    ensure_dir(OUTPUT_DIR)
    clean_output_for_run(
        raw_dir=raw_dir,
        md_dir=md_dir,
        master_md=master_md,
        docx_dir=docx_dir,
        summary_file=summary_file,
        skip_download=args.skip_download,
        skip_convert=args.skip_convert,
        skip_merge=args.skip_merge,
        skip_docx=args.skip_docx,
    )

    download_stats: DownloadStats | None = None
    convert_stats: ConvertStats | None = None
    merged_sources: int | None = None
    docx_files: list[Path] | None = None

    if not args.skip_download:
        print("[step] download")
        download_stats = download_all(manifest, raw_dir)

    if not args.skip_convert:
        print("[step] convert to markdown")
        convert_stats = convert_to_md(raw_dir, md_dir)

    if not args.skip_merge:
        print("[step] merge markdown")
        merged_sources = merge_md_files(md_dir, master_md)

    if not args.skip_docx:
        print("[step] build docx pack")
        if not master_md.exists():
            raise FileNotFoundError(
                f"Master Markdown file not found: {master_md}. Run the merge step first."
            )
        blocks = parse_master_md(master_md.read_text(encoding="utf-8", errors="ignore"))
        docx_files = build_docx_pack(
            blocks=blocks,
            out_dir=docx_dir,
            master_md=master_md,
            chunk_words=args.chunk_words,
            pack_words=args.pack_words,
            max_words=args.max_words,
            max_docx_parts=args.max_docx_parts,
        )

    write_summary(
        summary_path=summary_file,
        manifest=manifest,
        raw_dir=raw_dir,
        md_dir=md_dir,
        master_md=master_md,
        docx_dir=docx_dir,
        download_stats=download_stats,
        convert_stats=convert_stats,
        merged_sources=merged_sources,
        docx_files=docx_files,
    )
    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
